"""
DOCX Exporter - Export en format Microsoft Word
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
import sys
import os

# Ajouter le chemin parent pour importer i18n
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from core.i18n import I18n

class DOCXExporter:
    """Exporte le livre en DOCX"""
    
    def export(self, book_manager, output_dir: Path, lang='fr') -> Path:
        """
        Exporte le livre en DOCX
        
        Args:
            book_manager: Le gestionnaire de livre
            output_dir: Répertoire de sortie
            lang: Code langue ('fr', 'en', 'es', etc.)
        
        Returns:
            Path: Chemin du fichier créé
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        lang_suffix = f"_{lang.upper()}" if lang != 'fr' else ""
        filename = f"{book_manager.title.replace(' ', '_')}{lang_suffix}_{timestamp}.docx"
        filepath = output_dir / filename
        
        # Créer le document
        doc = Document()
        
        # Configurer les marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Page de titre
        title = doc.add_heading(book_manager.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Auteur avec traduction "Par" / "By" / etc.
        by_text = I18n.get('by', lang)
        author = doc.add_paragraph(f"{by_text} {book_manager.author}")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author.runs[0].font.size = Pt(14)
        author.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Espace
        
        # Sous-titre traduit
        subtitle_text = I18n.get('subtitle', lang)
        subtitle = doc.add_paragraph(subtitle_text)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(11)
        subtitle.runs[0].italic = True
        
        # Saut de page
        doc.add_page_break()
        
        # Chapitres
        for i, chapter in enumerate(book_manager.chapters):
            # Titre du chapitre traduit
            chapter_num = I18n.get_chapter_number(i+1, lang)
            chapter_title_text = chapter.get_title_translation(lang)
            chapter_title = doc.add_heading(f"{chapter_num}: {chapter_title_text}", level=1)
            chapter_title.runs[0].font.size = Pt(18)
            chapter_title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            # Contenu
            if lang == 'fr':
                content = chapter.content_fr
            else:
                content = chapter.get_translation(lang)
            
            if content and content.strip():
                # Diviser en paragraphes
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        para.runs[0].font.size = Pt(11)
            else:
                # Texte "chapitre vide" traduit
                empty_text = I18n.get('empty_chapter', lang)
                para = doc.add_paragraph(empty_text)
                para.runs[0].font.italic = True
                para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # Saut de page après chaque chapitre (sauf le dernier)
            if i < len(book_manager.chapters) - 1:
                doc.add_page_break()
        
        # Sauvegarder
        doc.save(str(filepath))
        
        return filepath

